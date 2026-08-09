"""fetch_document: parse PDF/DOCX/XLSX/CSV downloads into plain text.

Type is decided by MAGIC BYTES (first 2KB), never by the URL extension or
Content-Type: ``%PDF`` -> pdf; ``PK\\x03\\x04`` -> zip, then the entry list
decides docx (``word/``) vs xlsx (``xl/``); anything else is tried as CSV.

Honest statuses (phase-1 style, never fake success): ok /
unsupported_format / too_large / corrupted / timeout / extraction_failed /
ssrf_blocked / login_required / scan_rejected.

Parsing libraries (pypdf / pdfplumber / python-docx / openpyxl /
charset-normalizer) are lazy imports; the whole parse runs in a thread under
one overall timeout — pathological PDFs can otherwise loop forever.
"""

from __future__ import annotations

import asyncio
import csv
import io
import zipfile
from pathlib import PurePath
from typing import Protocol

from proseforge.infrastructure.webtools.fetcher import SafeFetcher

MAX_PDF_PAGES = 200
PARSE_TIMEOUT_SECONDS = 60.0
# Zip-bomb preflight (docx/xlsx are zips): reject before any real parsing.
MAX_ZIP_ENTRIES = 10_000
MAX_ZIP_RATIO = 100.0
MAX_ZIP_UNCOMPRESSED = 100 * 1024 * 1024
MAX_SHEET_ROWS = 10_000
MAX_SHEET_CELLS = 100_000
CSV_SAMPLE_BYTES = 64 * 1024
MAX_CSV_ROWS = 10_000
# pypdf extracting almost nothing means a scanned/image PDF.
MIN_EXTRACTED_CHARS = 20

STATUS_OK = "ok"
STATUS_UNSUPPORTED = "unsupported_format"
STATUS_TOO_LARGE = "too_large"
STATUS_CORRUPTED = "corrupted"
STATUS_TIMEOUT = "timeout"
STATUS_EXTRACTION_FAILED = "extraction_failed"
STATUS_SSRF_BLOCKED = "ssrf_blocked"
STATUS_LOGIN_REQUIRED = "login_required"
STATUS_SCAN_REJECTED = "scan_rejected"


class DocumentScanner(Protocol):
    """Malware-scan seam: return True when the bytes are safe to parse.

    Default is None (allow) — wire ClamAV here later without touching the
    call sites.
    """

    def scan(self, data: bytes) -> bool: ...


def _truncate_head_tail(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    omitted = len(text) - max_chars
    half = max_chars // 2
    return f"{text[:half]}\n[truncated: {omitted} chars]\n{text[-(max_chars - half):]}"


def sniff_document_kind(data: bytes) -> str:
    """pdf / docx / xlsx / csv from magic bytes (+ zip entries for OOXML)."""
    head = data[:2048]
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                names = archive.namelist()
        except zipfile.BadZipFile:
            return "corrupted"
        if any(name.startswith("word/") for name in names):
            return "docx"
        if any(name.startswith("xl/") for name in names):
            return "xlsx"
        return "unsupported_format"
    return "csv"


def _zip_preflight(data: bytes) -> str | None:
    """Reject zip bombs before parsing. Returns a status string or None."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
    except zipfile.BadZipFile:
        return STATUS_CORRUPTED
    if len(infos) > MAX_ZIP_ENTRIES:
        return STATUS_TOO_LARGE
    total_uncompressed = sum(info.file_size for info in infos)
    total_compressed = max(sum(info.compress_size for info in infos), 1)
    if total_uncompressed > MAX_ZIP_UNCOMPRESSED:
        return STATUS_TOO_LARGE
    if total_compressed > 1024 and total_uncompressed / total_compressed > MAX_ZIP_RATIO:
        return STATUS_TOO_LARGE
    return None


def _parse_pdf(data: bytes, max_length: int) -> dict:
    from pypdf import PdfReader  # lazy: api-extra dependency

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception:
        return {"status": STATUS_CORRUPTED, "kind": "pdf"}
    page_count = len(reader.pages)
    parts: list[str] = []
    for page in reader.pages[:MAX_PDF_PAGES]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")  # one broken page must not fail the document
        if sum(len(part) for part in parts) > max_length * 2:
            break
    text = "\n".join(part for part in parts if part).strip()
    if len(text) < MIN_EXTRACTED_CHARS:
        # Simple fallback: pdfplumber sometimes wins on table-heavy pages.
        text = _parse_pdf_pdfplumber(data, max_length)
    metadata = {"pages": page_count, "pages_processed": min(page_count, MAX_PDF_PAGES)}
    if len(text) < MIN_EXTRACTED_CHARS:
        # Scanned/image PDF: honest failure, no invented content.
        return {"status": STATUS_EXTRACTION_FAILED, "kind": "pdf", "text": "", "metadata": metadata, "note": "no extractable text (likely a scanned/image PDF)"}
    return {"status": STATUS_OK, "kind": "pdf", "text": text, "metadata": metadata}


def _parse_pdf_pdfplumber(data: bytes, max_length: int) -> str:
    try:
        import pdfplumber  # lazy: api-extra dependency
    except ImportError:
        return ""
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            parts = [(page.extract_text() or "") for page in pdf.pages[:MAX_PDF_PAGES]]
        return "\n".join(part for part in parts if part).strip()[: max_length * 2]
    except Exception:
        return ""


def _parse_docx(data: bytes) -> dict:
    rejected = _zip_preflight(data)
    if rejected is not None:
        return {"status": rejected, "kind": "docx"}
    import docx  # lazy: api-extra dependency (python-docx)

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        return {"status": STATUS_CORRUPTED, "kind": "docx"}
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n".join(paragraphs).strip()
    metadata = {"paragraphs": len(document.paragraphs)}
    if len(text) < MIN_EXTRACTED_CHARS:
        return {"status": STATUS_EXTRACTION_FAILED, "kind": "docx", "text": "", "metadata": metadata, "note": "no extractable text"}
    return {"status": STATUS_OK, "kind": "docx", "text": text, "metadata": metadata}


def _parse_xlsx(data: bytes) -> dict:
    rejected = _zip_preflight(data)
    if rejected is not None:
        return {"status": rejected, "kind": "xlsx"}
    import openpyxl  # lazy: api-extra dependency

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception:
        return {"status": STATUS_CORRUPTED, "kind": "xlsx"}
    lines: list[str] = []
    total_rows = 0
    total_cells = 0
    sheet_names = list(workbook.sheetnames)
    truncated = False
    for sheet in workbook.worksheets:
        lines.append(f"# sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            total_rows += 1
            total_cells += len(row)
            if total_rows > MAX_SHEET_ROWS or total_cells > MAX_SHEET_CELLS:
                truncated = True
                break
            values = ["" if cell is None else str(cell) for cell in row]
            lines.append("\t".join(values))
        if truncated:
            break
    workbook.close()
    text = "\n".join(lines).strip()
    metadata = {"sheets": len(sheet_names), "sheet_names": sheet_names, "rows": total_rows, "cells": total_cells, "rows_truncated": truncated}
    if not text:
        return {"status": STATUS_EXTRACTION_FAILED, "kind": "xlsx", "text": "", "metadata": metadata, "note": "workbook is empty"}
    return {"status": STATUS_OK, "kind": "xlsx", "text": text, "metadata": metadata}


def _parse_csv(data: bytes) -> dict:
    sample = data[:CSV_SAMPLE_BYTES]
    encoding = "utf-8-sig" if sample.startswith(b"\xef\xbb\xbf") else _detect_encoding(sample)
    try:
        text = data.decode(encoding, errors="replace")
    except (LookupError, ValueError):
        text = data.decode("utf-8", errors="replace")
        encoding = "utf-8"
    sample_text = text[:CSV_SAMPLE_BYTES]
    try:
        dialect = csv.Sniffer().sniff(sample_text, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel  # Sniffer is a heuristic; comma is the safe default
    rows: list[list[str]] = []
    truncated = False
    try:
        for row in csv.reader(io.StringIO(text), dialect):
            rows.append(row)
            if len(rows) >= MAX_CSV_ROWS:
                truncated = True
                break
    except csv.Error:
        return {"status": STATUS_CORRUPTED, "kind": "csv"}
    rendered = "\n".join(dialect.delimiter.join(row) for row in rows).strip()
    metadata = {"encoding": encoding, "delimiter": dialect.delimiter, "rows": len(rows), "rows_truncated": truncated}
    if not rendered:
        return {"status": STATUS_EXTRACTION_FAILED, "kind": "csv", "text": "", "metadata": metadata, "note": "no readable rows"}
    return {"status": STATUS_OK, "kind": "csv", "text": rendered, "metadata": metadata}


def _detect_encoding(sample: bytes) -> str:
    try:
        from charset_normalizer import from_bytes  # lazy: api-extra dependency
    except ImportError:
        return "utf-8"
    best = from_bytes(sample).best()
    return str(best.encoding) if best is not None else "utf-8"


_PARSERS = {"pdf": _parse_pdf, "docx": _parse_docx, "xlsx": _parse_xlsx, "csv": _parse_csv}

# Chat-attachment injection whitelist (M1): text formats decode directly,
# PDF/OOXML go through the magic-byte sniff + parsers above. Images are
# deliberately out of scope.
TEXT_PARSE_EXTENSIONS = frozenset({".txt", ".md", ".json", ".csv"})
BINARY_PARSE_EXTENSIONS = frozenset({".pdf", ".docx", ".xlsx"})
SUPPORTED_PARSE_EXTENSIONS = TEXT_PARSE_EXTENSIONS | BINARY_PARSE_EXTENSIONS
# Character ceiling for one parsed upload: bounds parser output held in
# memory (token-level trimming is the injection layer's job).
UPLOAD_PARSE_MAX_CHARS = 200_000


def parse_document_bytes(data: bytes, filename: str) -> str:
    """Parse uploaded file bytes into plain text for context injection.

    txt/md/json/csv decode as UTF-8 (BOM tolerated, undecodable bytes
    replaced); pdf/docx/xlsx are decided by magic bytes (never the
    extension) and parsed with the same parsers fetch_document uses. Raises
    ValueError for unsupported extensions, unparseable content and honest
    extraction failures — the caller decides how to degrade.
    """
    suffix = PurePath(filename).suffix.lower()
    if suffix in TEXT_PARSE_EXTENSIONS:
        return data.decode("utf-8-sig", errors="replace").strip()
    if suffix not in BINARY_PARSE_EXTENSIONS:
        raise ValueError(f"unsupported file type: {suffix or filename!r}")
    kind = sniff_document_kind(data)
    if kind not in _PARSERS:
        raise ValueError(f"file content does not parse as {suffix[1:]} ({kind})")
    parser = _PARSERS[kind]
    result = parser(data, UPLOAD_PARSE_MAX_CHARS) if kind == "pdf" else parser(data)
    if result.get("status") != STATUS_OK:
        raise ValueError(f"could not extract text from {filename}: {result.get('note') or result.get('status')}")
    return str(result.get("text", "")).strip()[:UPLOAD_PARSE_MAX_CHARS]


async def fetch_document(url: str, *, max_length: int = 8000, fetcher: SafeFetcher, scanner: DocumentScanner | None = None) -> dict:
    """Download and parse a document URL into plain text. Never raises for
    expected failures — everything maps to an honest status."""
    outcome = await fetcher.fetch_bytes(url)
    if outcome.too_large:
        return {"status": STATUS_TOO_LARGE, "url": url, "error": outcome.error}
    if outcome.error is not None or outcome.data is None:
        if outcome.error_kind == "ssrf":
            status = STATUS_SSRF_BLOCKED
        elif outcome.error_kind == "timeout":
            status = STATUS_TIMEOUT
        elif outcome.status_code in (401, 403):
            status = STATUS_LOGIN_REQUIRED
        else:
            status = STATUS_EXTRACTION_FAILED
        return {"status": status, "url": url, "error": outcome.error or "empty response"}
    data = outcome.data
    if scanner is not None and not scanner.scan(data):
        return {"status": STATUS_SCAN_REJECTED, "url": outcome.final_url}
    kind = sniff_document_kind(data)
    if kind == "unsupported_format":
        return {"status": STATUS_UNSUPPORTED, "url": outcome.final_url, "note": "zip container is neither docx nor xlsx"}
    if kind == "corrupted":
        return {"status": STATUS_CORRUPTED, "url": outcome.final_url}
    parser = _PARSERS[kind]
    try:
        if kind == "pdf":
            result = await asyncio.wait_for(asyncio.to_thread(parser, data, max_length * 2), timeout=PARSE_TIMEOUT_SECONDS)
        else:
            result = await asyncio.wait_for(asyncio.to_thread(parser, data), timeout=PARSE_TIMEOUT_SECONDS)
    except TimeoutError:
        return {"status": STATUS_TIMEOUT, "kind": kind, "url": outcome.final_url, "note": f"parse exceeded {PARSE_TIMEOUT_SECONDS:.0f}s"}
    except Exception:
        return {"status": STATUS_CORRUPTED, "kind": kind, "url": outcome.final_url}
    result["text"] = _truncate_head_tail(str(result.get("text", "")), max_length)
    result["url"] = outcome.final_url
    return result
